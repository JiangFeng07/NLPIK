#!/usr/local/bin/ python3
# -*- coding: utf-8 -*-
# @Time : 2022-04-01 10:49 
# @Author : Leo

import docx

file = docx.Document('/tmp/doc/test13.docx')
file2 = docx.Document('/tmp/doc/test1.docx')

doc = docx.Document()
text_list=[]
for para in file:
    text_list.append(para.text)
print(''.join(text_list))
#     doc.add_paragraph(para.text)
# for para in file.paragraphs:
#     doc.add_paragraph(para.text)
# doc.save('/tmp/doc/merge.docx')
